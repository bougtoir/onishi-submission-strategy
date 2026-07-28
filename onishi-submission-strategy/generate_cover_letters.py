#!/usr/bin/env python3
"""Generate cover letters for 4 ONISHI framework papers."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Paragraph spacing
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(8)
    doc.add_page_break()


def add_header(doc, label, color_rgb):
    """Add a colored header label for each letter."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f'Cover Letter — {label}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color_rgb


def add_date_and_address(doc, editor_info):
    """Add date and editor address block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('[Date]')
    run.font.color.rgb = RGBColor(180, 180, 180)

    for line in editor_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.add_run(line)


def add_body(doc, paragraphs):
    """Add body paragraphs."""
    for para_text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        # Handle bold markers **text**
        parts = para_text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True


def add_closing(doc):
    """Add closing and signature block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run('Sincerely,')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('Tatsuki Onishi')
    run.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    p2.add_run('Shiga University')
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    run3 = p3.add_run('[email address]')
    run3.font.color.rgb = RGBColor(180, 180, 180)


# ============================================================
# 1. IONE → BMC MRM Special Issue
# ============================================================
add_header(doc, 'IONE → BMC Medical Research Methodology (Special Issue)', RGBColor(0xFF, 0x98, 0x00))

add_date_and_address(doc, [
    'Guest Editors',
    'Dr. Rishi J. Desai, Dr. Ivan Olier, Dr. Joy Shi',
    'Special Issue: "Causal inference and observational data vol. 2"',
    'BMC Medical Research Methodology',
])

add_body(doc, [
    'Dear Drs. Desai, Olier, and Shi,',

    'We are pleased to submit our manuscript entitled **"IONE: Incoherence-Oriented Neutralization and Extraction — '
    'A Framework for Detecting Hidden Subgroup Structures in Cohort Studies"** for consideration in the special issue '
    '"Causal inference and observational data vol. 2" of BMC Medical Research Methodology.',

    'A fundamental but often overlooked challenge in causal inference from observational data is the presence of '
    'latent population heterogeneity — hidden subgroups whose treatment responses differ qualitatively from the '
    'overall cohort. When such incoherent subpopulations remain undetected, standard analytical methods may yield '
    'effect estimates that represent no actual individual within the study population, leading to potentially '
    'misleading causal conclusions.',

    'The IONE framework addresses this challenge by introducing the **C1 coherence index**, a quantitative measure '
    'that evaluates whether the observed treatment effect in a cohort study faithfully represents its constituent '
    'subpopulations. When incoherence is detected, IONE provides a systematic decomposition procedure to identify '
    'and characterize these hidden subgroups, enabling more accurate causal inference at the subpopulation level.',

    'This work is directly relevant to the special issue\'s focus on causal inference methods for observational data. '
    'Specifically, IONE contributes a methodological tool that strengthens the internal validity of cohort-based '
    'causal analyses by ensuring that population-level estimates are not artifacts of unrecognized heterogeneity. '
    'We believe this addresses a critical gap between study design assumptions and the analytical reality of '
    'observational datasets.',

    'The manuscript has not been published elsewhere and is not under consideration by another journal. '
    'A preprint version is available on medRxiv [DOI to be inserted]. '
    'All authors have approved the manuscript and agree with its submission to this journal.',

    'We confirm that this work complies with BMC Medical Research Methodology\'s editorial policies, '
    'including open data and ethical standards. We have no conflicts of interest to declare.',

    'Thank you for your consideration. We look forward to your response.',
])

add_closing(doc)
add_separator(doc)

# ============================================================
# 2. KOTHA → JCE
# ============================================================
add_header(doc, 'KOTHA → Journal of Clinical Epidemiology', RGBColor(0x4C, 0xAF, 0x50))

add_date_and_address(doc, [
    'Editor-in-Chief',
    'Journal of Clinical Epidemiology',
])

add_body(doc, [
    'Dear Editor,',

    'We are pleased to submit our manuscript entitled **"KOTHA: Knowledge-driven Observational-Trial '
    'Harmonization Approach — A Three-Module Framework for Diagnosing Structural Information Loss in '
    'RCT Meta-Analyses"** for consideration in the Journal of Clinical Epidemiology.',

    'Meta-analyses of randomized controlled trials (RCTs) are widely regarded as the highest level of '
    'evidence, yet they can suffer from structural information loss when treatment effects are heterogeneous '
    'across trials. Current methods for assessing heterogeneity (e.g., I-squared, tau-squared) quantify '
    'statistical inconsistency but do not diagnose **why** information is lost or **how** observational data '
    'might restore it.',

    'KOTHA addresses this gap through three integrated modules: '
    '**Module K** (counterfactual power simulation) identifies which RCTs contribute meaningful information '
    'and which introduce structural noise; '
    '**Module T** (hierarchical Bayesian integration) synthesizes RCT and observational evidence with '
    'appropriate uncertainty propagation; and '
    '**Module H** (GRADE extension) provides a structured framework for transparently rating the quality '
    'of the harmonized evidence.',

    'We believe this manuscript aligns closely with JCE\'s mission to publish innovative methodological '
    'developments that directly improve clinical evidence synthesis. The GRADE extension component '
    '(Module H) is particularly relevant given JCE\'s central role in advancing the GRADE framework. '
    'Our approach provides clinical epidemiologists with practical, actionable tools for situations where '
    'RCT evidence alone may be insufficient or misleading.',

    'The manuscript has not been published elsewhere and is not under consideration by another journal. '
    'A preprint version is available on medRxiv [DOI to be inserted]. '
    'All authors have approved the manuscript and agree with its submission to this journal.',

    'We have no conflicts of interest to declare.',

    'Thank you for your consideration.',
])

add_closing(doc)
add_separator(doc)

# ============================================================
# 3. LINKO → RSM
# ============================================================
add_header(doc, 'LINKO → Research Synthesis Methods', RGBColor(0x21, 0x96, 0xF3))

add_date_and_address(doc, [
    'Editor-in-Chief',
    'Research Synthesis Methods',
])

add_body(doc, [
    'Dear Editor,',

    'We are pleased to submit our manuscript entitled **"LINKO: Latent Information Normalization for '
    'Key Outcomes — Evaluating Meta-Analysis Validity through Information Contribution Ratio"** '
    'for consideration in Research Synthesis Methods.',

    'A common but underappreciated problem in meta-analysis is that individual studies contribute '
    'information of vastly different quality and relevance to the pooled estimate, yet conventional '
    'weighting schemes (inverse-variance, sample size) do not capture this multidimensional notion '
    'of "information contribution." As a result, a meta-analysis may be dominated by studies that '
    'contribute statistical precision but limited inferential value.',

    'LINKO introduces the **Information Contribution Ratio (ICR)**, a novel metric that quantifies '
    'each study\'s true informational contribution to the meta-analytic estimate by integrating '
    'precision, relevance, and methodological quality into a single normalized score. We also '
    'introduce the **Prism Forest Plot**, a new visualization that overlays ICR values onto the '
    'traditional forest plot, enabling researchers to immediately identify which studies are driving '
    'the pooled estimate and whether that influence is justified.',

    'This work falls squarely within RSM\'s scope as a methodological contribution to research '
    'synthesis. The ICR metric provides a practical tool that meta-analysts can apply alongside '
    'existing methods (GRADE, risk of bias assessment) to evaluate and communicate the robustness '
    'of their findings. The Prism Forest Plot offers an intuitive visual complement to standard '
    'reporting.',

    'The manuscript has not been published elsewhere and is not under consideration by another journal. '
    'A preprint version is available on medRxiv [DOI to be inserted]. '
    'All authors have approved the manuscript and agree with its submission to this journal.',

    'We have no conflicts of interest to declare.',

    'Thank you for your consideration.',
])

add_closing(doc)
add_separator(doc)

# ============================================================
# 4. ONISHI → AJE
# ============================================================
add_header(doc, 'ONISHI → American Journal of Epidemiology', RGBColor(0x9C, 0x27, 0xB0))

add_date_and_address(doc, [
    'Editor-in-Chief',
    'American Journal of Epidemiology',
])

add_body(doc, [
    'Dear Editor,',

    'We are pleased to submit our manuscript entitled **"ONISHI: Optimal Normalization, Incoherence '
    'Stratification, and Harmonized Integration — A Unified Framework for Multi-level Evidence Assessment"** '
    'for consideration in the American Journal of Epidemiology.',

    'Epidemiological evidence synthesis increasingly requires researchers to navigate three interrelated '
    'challenges simultaneously: evaluating the information quality of individual studies within a '
    'meta-analysis, detecting hidden population heterogeneity in cohort data, and harmonizing evidence '
    'across randomized and observational study designs. Existing methods address these challenges in '
    'isolation, leaving researchers without a coherent analytical pipeline for complex evidence landscapes.',

    'The ONISHI framework unifies three complementary methodologies into an integrated evidence assessment '
    'pipeline: '
    '**LINKO** (information contribution quantification via the ICR metric) normalizes the informational '
    'value of each study; '
    '**IONE** (incoherence detection and subgroup decomposition via the C1 index) identifies hidden '
    'population structures that may confound aggregate estimates; and '
    '**KOTHA** (RCT-observational harmonization via counterfactual simulation and hierarchical Bayesian '
    'synthesis) bridges the gap between experimental and real-world evidence.',

    'We demonstrate that the sequential or simultaneous application of these three components yields '
    'analytical capabilities that are not achievable by any single method alone, including '
    'subpopulation-level information contribution analysis, structure-aware Bayesian evidence integration, '
    'and population-decomposed external validity assessment. We illustrate these capabilities through '
    'clinical scenarios in diabetes management, statin therapy, and magnesium supplementation.',

    'This work contributes to AJE\'s scope of methodological developments in epidemiology by providing '
    'a practical, modular framework that researchers can adopt incrementally. Each component can be used '
    'independently or in combination, making it accessible to epidemiologists at varying levels of '
    'methodological sophistication.',

    'The three component methodologies have been described in detail in separate manuscripts: '
    'LINKO (under review at Research Synthesis Methods), '
    'IONE (submitted to BMC Medical Research Methodology, special issue on causal inference), and '
    'KOTHA (under review at Journal of Clinical Epidemiology). '
    'Preprint versions of all three are available on medRxiv [DOIs to be inserted]. '
    'The present manuscript focuses on the integration of these methods and the novel capabilities '
    'that emerge from their combination.',

    'The manuscript has not been published elsewhere and is not under consideration by another journal. '
    'All authors have approved the manuscript and agree with its submission to this journal.',

    'We have no conflicts of interest to declare.',

    'Thank you for your consideration.',
])

add_closing(doc)

# Save
output_path = '/home/ubuntu/cover_letters_ONISHI_4papers.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
