#!/usr/bin/env python3
"""
LINKO × IONE × KOTHA 組み合わせ報告書
カラー図版入り編集可能 .docx ファイル生成
"""

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import os

# Japanese font setup
plt.rcParams['font.family'] = 'DejaVu Sans'

OUT_DIR = "/home/ubuntu/report_figures"
os.makedirs(OUT_DIR, exist_ok=True)

# =========================================================
# Color palette
# =========================================================
C_LINKO = "#2196F3"   # Blue
C_IONE  = "#FF9800"   # Orange
C_KOTHA = "#4CAF50"   # Green
C_ALL3  = "#9C27B0"   # Purple (all three combined)
C_BG    = "#FAFAFA"
C_ARROW = "#616161"

def make_rounded_box(ax, x, y, w, h, text, color, fontsize=11, text_color="white"):
    """Draw a rounded rectangle with centered text."""
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02",
                         facecolor=color, edgecolor="white", linewidth=2, alpha=0.92)
    ax.add_patch(box)
    ax.text(x + w/2, y + h/2, text, ha="center", va="center",
            fontsize=fontsize, fontweight="bold", color=text_color,
            wrap=True)

# =========================================================
# Figure 1: Three Frameworks Overview
# =========================================================
def fig1_overview():
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Title
    ax.text(6, 6.6, "Three Analytical Frameworks: LINKO, IONE, KOTHA",
            ha="center", va="center", fontsize=16, fontweight="bold", color="#333")

    # LINKO box
    make_rounded_box(ax, 0.3, 3.15, 3.4, 2.95, "", C_LINKO)
    ax.text(2.0, 5.78, "LINKO", ha="center", va="center", fontsize=14,
            fontweight="bold", color="white")
    ax.text(2.0, 5.35, "Latent Information Normalization\nfor Key Outcomes",
            ha="center", va="center", fontsize=8.5, color="#EEEEEE")
    lines_linko = [
        "ICR (information contribution ratio)",
        "Meta-analytic validity diagnostic",
        "Prism forest plot",
        "ICR-weighted pooling",
    ]
    for i, line in enumerate(lines_linko):
        ax.text(2.0, 4.72 - i*0.33, f"\u2022 {line}", ha="center", va="center",
                fontsize=8, color="white")

    # IONE box
    make_rounded_box(ax, 4.3, 3.15, 3.4, 2.95, "", C_IONE)
    ax.text(6.0, 5.78, "IONE", ha="center", va="center", fontsize=14,
            fontweight="bold", color="white")
    ax.text(6.0, 5.35, "Incoherence-Oriented Neutralisation\nand Extraction",
            ha="center", va="center", fontsize=8.5, color="#EEEEEE")
    lines_ione = [
        "Coherent subgroup detection",
        "C1 incoherence indicator",
        "Within-stratum homogeneity (W)",
        "Hidden structure identification",
    ]
    for i, line in enumerate(lines_ione):
        ax.text(6.0, 4.72 - i*0.33, f"\u2022 {line}", ha="center", va="center",
                fontsize=8, color="white")

    # KOTHA box
    make_rounded_box(ax, 8.3, 3.15, 3.4, 2.95, "", C_KOTHA)
    ax.text(10.0, 5.78, "KOTHA", ha="center", va="center", fontsize=14,
            fontweight="bold", color="white")
    ax.text(10.0, 5.35, "Knowledge-driven Observational-Trial\nHarmonisation Approach",
            ha="center", va="center", fontsize=8.5, color="#EEEEEE")
    lines_kotha = [
        "Module K: counterfactual power",
        "Module T: Bayesian integration",
        "Module H: OIS / TSA / GRADE",
        "RCT-observational harmonisation",
    ]
    for i, line in enumerate(lines_kotha):
        ax.text(10.0, 4.72 - i*0.33, f"\u2022 {line}", ha="center", va="center",
                fontsize=8, color="white")

    # Target level labels
    ax.text(2.0, 2.85, "Level: between studies",
            ha="center", fontsize=8, color=C_LINKO, fontstyle="italic")
    ax.text(6.0, 2.85, "Level: within studies",
            ha="center", fontsize=8, color=C_IONE, fontstyle="italic")
    ax.text(10.0, 2.85, "Level: between study types",
            ha="center", fontsize=8, color=C_KOTHA, fontstyle="italic")

    # Common data layer
    make_rounded_box(ax, 1.5, 0.5, 9.0, 1.5, "", "#78909C")
    ax.text(6.0, 1.55, "Common Data Sources", ha="center", va="center",
            fontsize=12, fontweight="bold", color="white")
    ax.text(6.0, 1.0, "RCT published data (Table 1)  |  Individual Patient Data (IPD)  |  Observational cohort data",
            ha="center", va="center", fontsize=9, color="white")

    # Arrows from data to frameworks
    for x_pos in [2.0, 6.0, 10.0]:
        ax.annotate("", xy=(x_pos, 3.15), xytext=(x_pos, 2.0),
                    arrowprops=dict(arrowstyle="-|>", color=C_ARROW, lw=1.5))

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig1_overview.png")
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

# =========================================================
# Figure 2: Four Combinations
# =========================================================
def fig2_four_combinations():
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    fig.patch.set_facecolor("white")
    fig.suptitle("4 Combinations of LINKO, IONE, and KOTHA",
                 fontsize=16, fontweight="bold", color="#333", y=0.98)

    combos = [
        ("LINKO + IONE", [C_LINKO, C_IONE],
         "ICR-guided Subpopulation Analysis",
         ["LINKO identifies anomalous ICR\nacross studies in meta-analysis",
          "IONE decomposes each study\ninto coherent subgroups",
          "ICR recalculated per subgroup\nreveals hidden heterogeneity sources",
          "Subgroup-specific ICR-weighted\nmeta-analysis"]),
        ("LINKO + KOTHA", [C_LINKO, C_KOTHA],
         "Information-Aware Evidence Synthesis",
         ["LINKO quantifies information\ncontribution of endpoints",
          "ICR feeds into KOTHA Module K\nfor refined power simulation",
          "Module T uses ICR as Bayesian\nprior weight for integration",
          "Module H reports ICR alongside\nOIS and TSA assessments"]),
        ("IONE + KOTHA", [C_IONE, C_KOTHA],
         "Population-Harmonized Evidence Integration",
         ["IONE detects incoherent\nsubgroups in observational data",
          "Subgroup risk profiles inform\nKOTHA Module K simulation",
          "Module T integrates subgroup-\nspecific effects with RCT data",
          "Module H assesses representativeness\nusing IONE's C1 index"]),
        ("LINKO + IONE + KOTHA", [C_LINKO, C_IONE, C_KOTHA],
         "Comprehensive Evidence Assessment Pipeline",
         ["IONE decomposes populations\ninto coherent subgroups",
          "LINKO quantifies ICR at\nsubgroup level per study",
          "KOTHA integrates all evidence\nwith ICR + C1 weighting",
          "Full pipeline: detect heterogeneity\n-> quantify info -> harmonize"]),
    ]

    for idx, (title, colors, subtitle, steps) in enumerate(combos):
        ax = axes[idx // 2][idx % 2]
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 8)
        ax.axis("off")

        # Title bar
        if len(colors) == 2:
            # Gradient-like effect with two colors
            rect1 = FancyBboxPatch((0, 7), 5, 0.8, boxstyle="round,pad=0.02",
                                    facecolor=colors[0], alpha=0.85)
            rect2 = FancyBboxPatch((5, 7), 5, 0.8, boxstyle="round,pad=0.02",
                                    facecolor=colors[1], alpha=0.85)
            ax.add_patch(rect1)
            ax.add_patch(rect2)
        else:
            # Three colors
            for i, c in enumerate(colors):
                rect = FancyBboxPatch((i*3.33, 7), 3.34, 0.8,
                                       boxstyle="round,pad=0.02",
                                       facecolor=c, alpha=0.85)
                ax.add_patch(rect)

        ax.text(5, 7.4, title, ha="center", va="center",
                fontsize=13, fontweight="bold", color="white")

        # Subtitle
        ax.text(5, 6.55, subtitle, ha="center", va="center",
                fontsize=10, color="#444", fontstyle="italic")

        # Steps as flow
        step_colors = ["#E3F2FD", "#FFF3E0", "#E8F5E9", "#F3E5F5"]
        border_colors = [C_LINKO, C_IONE, C_KOTHA, C_ALL3]
        y_positions = [5.3, 3.8, 2.3, 0.8]

        for i, (step, yp) in enumerate(zip(steps, y_positions)):
            box = FancyBboxPatch((0.5, yp), 9.0, 1.2,
                                  boxstyle="round,pad=0.02",
                                  facecolor=step_colors[i],
                                  edgecolor=border_colors[i] if idx < 3 else colors[min(i, len(colors)-1)],
                                  linewidth=2)
            ax.add_patch(box)
            ax.text(0.9, yp + 0.85, f"Step {i+1}", fontsize=8,
                    fontweight="bold", color=border_colors[i] if idx < 3 else colors[min(i, len(colors)-1)])
            ax.text(5.0, yp + 0.45, step, ha="center", va="center",
                    fontsize=9, color="#333")

            # Arrow between steps
            if i < 3:
                ax.annotate("", xy=(5, yp + 1.2), xytext=(5, yp + 1.5),
                            arrowprops=dict(arrowstyle="-|>", color=C_ARROW, lw=1.2))

    plt.tight_layout(rect=[0, 0, 1, 0.96])
    path = os.path.join(OUT_DIR, "fig2_four_combinations.png")
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

# =========================================================
# Figure 3: Synergy Matrix
# =========================================================
def fig3_synergy_matrix():
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(5, 7.6, "Synergy Matrix: What Each Combination Enables",
            ha="center", va="center", fontsize=15, fontweight="bold", color="#333")

    # Table header
    headers = ["Combination", "New Capability", "Gain vs. Solo"]
    col_x = [0.3, 3.0, 7.5]
    col_w = [2.5, 4.2, 2.3]

    for i, (h, x, w) in enumerate(zip(headers, col_x, col_w)):
        rect = FancyBboxPatch((x, 6.8), w, 0.6, boxstyle="round,pad=0.01",
                               facecolor="#37474F", edgecolor="white", linewidth=1)
        ax.add_patch(rect)
        ax.text(x + w/2, 7.1, h, ha="center", va="center",
                fontsize=10, fontweight="bold", color="white")

    rows = [
        ("LINKO\n+\nIONE", C_LINKO, C_IONE,
         "Subgroup-level ICR analysis\nreveals why meta-analysis\nheterogeneity arises",
         "Root cause\nidentification"),
        ("LINKO\n+\nKOTHA", C_LINKO, C_KOTHA,
         "ICR-informed power simulation\nand Bayesian weighting improve\nevidence synthesis precision",
         "Quantitative\nprecision"),
        ("IONE\n+\nKOTHA", C_IONE, C_KOTHA,
         "Population decomposition feeds\ncounterfactual simulation and\nrepresentativeness assessment",
         "External\nvalidity"),
        ("LINKO + IONE\n+\nKOTHA", C_ALL3, C_ALL3,
         "Full pipeline: detect heterogeneity\n-> quantify information ->\nharmonize all evidence",
         "Comprehensive\nassessment"),
    ]

    for i, (combo, c1, c2, capability, gain) in enumerate(rows):
        y = 5.2 - i * 1.5
        # Row bg
        bg_color = "#F5F5F5" if i % 2 == 0 else "white"
        rect = FancyBboxPatch((0.3, y), 9.5, 1.3, boxstyle="round,pad=0.01",
                               facecolor=bg_color, edgecolor="#E0E0E0", linewidth=1)
        ax.add_patch(rect)

        # Combo name with colored indicator
        ax.text(1.55, y + 0.65, combo, ha="center", va="center",
                fontsize=9, fontweight="bold", color="#333")
        # Color dots
        if c1 != c2:
            ax.plot(0.5, y + 0.85, 'o', color=c1, markersize=8)
            ax.plot(0.5, y + 0.45, 'o', color=c2, markersize=8)
        else:
            ax.plot(0.5, y + 0.95, 'o', color=C_LINKO, markersize=6)
            ax.plot(0.5, y + 0.65, 'o', color=C_IONE, markersize=6)
            ax.plot(0.5, y + 0.35, 'o', color=C_KOTHA, markersize=6)

        # Capability
        ax.text(5.1, y + 0.65, capability, ha="center", va="center",
                fontsize=9, color="#333")

        # Gain
        ax.text(8.65, y + 0.65, gain, ha="center", va="center",
                fontsize=9, fontweight="bold", color=c1)

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig3_synergy_matrix.png")
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

# =========================================================
# Figure 4: Integrated Pipeline Diagram
# =========================================================
def fig4_pipeline():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(7, 7.6, "Integrated Pipeline: LINKO + IONE + KOTHA",
            ha="center", va="center", fontsize=16, fontweight="bold", color="#333")

    # Phase boxes
    phases = [
        (0.3, 5.35, 3.0, 2.05, "Phase 1: IONE\nPopulation decomposition", C_IONE,
         ["Detect incoherent populations", "Extract coherent subgroups", "Compute C1 and W"]),
        (4.0, 5.35, 3.0, 2.05, "Phase 2: LINKO\nInformation quantification", C_LINKO,
         ["Compute ICR per subgroup", "Assess ICR discrepancy", "Generate prism forest plot"]),
        (7.7, 5.35, 3.0, 2.05, "Phase 3: KOTHA\nEvidence harmonisation", C_KOTHA,
         ["Module K: power simulation", "Module T: Bayesian synthesis", "Module H: OIS / TSA / GRADE"]),
        (11.4, 5.35, 2.3, 2.05, "Output:\nDecision support", C_ALL3,
         ["Integrated evidence profile", "Subgroup-specific readout", "Uncertainty quantification"]),
    ]

    for x, y, w, h, title, color, items in phases:
        make_rounded_box(ax, x, y, w, h, "", color)
        ax.text(x + w/2, y + h - 0.28, title.split("\n")[0],
                ha="center", va="center", fontsize=9, fontweight="bold", color="white")
        if "\n" in title:
            ax.text(x + w/2, y + h - 0.62, title.split("\n")[1],
                    ha="center", va="center", fontsize=8, color="#F0F0F0")
        for i, item in enumerate(items):
            ax.text(x + w/2, y + h - 1.10 - i*0.32, f"\u2022 {item}",
                    ha="center", va="center", fontsize=7.5, color="white")

    # Arrows between phases
    arrow_positions = [(3.3, 6.375, 4.0, 6.375),
                       (7.0, 6.375, 7.7, 6.375),
                       (10.7, 6.375, 11.4, 6.375)]
    for x1, y1, x2, y2 in arrow_positions:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=C_ARROW, lw=2))

    # Data input layer
    make_rounded_box(ax, 0.3, 0.3, 13.4, 1.5, "", "#546E7A")
    ax.text(7, 1.45, "Input Data Sources", ha="center", va="center",
            fontsize=12, fontweight="bold", color="white")
    data_items = [
        ("Observational\nCohort Data\n(for IONE)", 2.5),
        ("Published RCT\nTable 1 Data\n(for LINKO)", 5.5),
        ("Individual Patient\nData (IPD)\n(for LINKO + IONE)", 8.5),
        ("RCT + Obs. Study\nEffect Estimates\n(for KOTHA)", 11.5),
    ]
    for label, x in data_items:
        ax.text(x, 0.75, label, ha="center", va="center",
                fontsize=8, color="white")

    # Feedback arrows
    # IONE -> KOTHA Module K
    ax.annotate("", xy=(9.2, 5.35), xytext=(1.8, 5.35),
                arrowprops=dict(arrowstyle="-|>", color=C_IONE,
                               connectionstyle="arc3,rad=-0.45", lw=1.5, ls="--"))
    ax.text(4.6, 3.55, "Subgroup risk profiles\nfeed Module K simulation",
            ha="center", va="center", fontsize=7.5, color=C_IONE,
            fontstyle="italic",
            bbox=dict(boxstyle="round,pad=0.25", facecolor="#FFF3E0", edgecolor=C_IONE, alpha=0.9))

    # LINKO -> KOTHA Module T
    ax.annotate("", xy=(9.2, 5.55), xytext=(5.5, 5.35),
                arrowprops=dict(arrowstyle="-|>", color=C_LINKO,
                               connectionstyle="arc3,rad=-0.35", lw=1.5, ls="--"))
    ax.text(8.7, 4.25, "ICR as Bayesian weight\nin Module T",
            ha="center", va="center", fontsize=7.5, color=C_LINKO,
            fontstyle="italic",
            bbox=dict(boxstyle="round,pad=0.25", facecolor="#E3F2FD", edgecolor=C_LINKO, alpha=0.9))

    # Data arrows up
    for x in [1.8, 5.5, 9.2, 12.55]:
        ax.annotate("", xy=(x, 5.35), xytext=(x, 1.8),
                    arrowprops=dict(arrowstyle="-|>", color="#90A4AE", lw=1))

    plt.tight_layout()
    path = os.path.join(OUT_DIR, "fig4_pipeline.png")
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

# =========================================================
# Figure 5: Combination comparison radar chart
# =========================================================
def fig5_radar():
    categories = [
        "Heterogeneity\nDiagnosis",
        "Population\nDecomposition",
        "Information\nQuantification",
        "Evidence\nIntegration",
        "Guideline\nUtility",
        "External\nValidity",
    ]
    N = len(categories)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]

    combos = {
        "LINKO + IONE":       [5, 4, 5, 2, 2, 3],
        "LINKO + KOTHA":      [4, 1, 5, 5, 4, 3],
        "IONE + KOTHA":       [3, 5, 2, 4, 4, 5],
        "LINKO+IONE+KOTHA":   [5, 5, 5, 5, 5, 5],
    }

    colors_map = {
        "LINKO + IONE": "#1565C0",
        "LINKO + KOTHA": "#2E7D32",
        "IONE + KOTHA": "#E65100",
        "LINKO+IONE+KOTHA": "#6A1B9A",
    }

    fig, ax = plt.subplots(figsize=(9, 9), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor("white")

    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_rlabel_position(0)

    plt.xticks(angles[:-1], categories, fontsize=10, fontweight="bold")
    plt.yticks([1, 2, 3, 4, 5], ["1", "2", "3", "4", "5"],
               color="grey", size=8)
    plt.ylim(0, 5.5)

    for name, values in combos.items():
        vals = values + values[:1]
        ax.plot(angles, vals, 'o-', linewidth=2, label=name,
                color=colors_map[name], markersize=6)
        ax.fill(angles, vals, alpha=0.08, color=colors_map[name])

    ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.1), fontsize=10,
              framealpha=0.9)
    ax.set_title("Capability Profile by Combination",
                 size=14, fontweight="bold", color="#333", pad=20)

    path = os.path.join(OUT_DIR, "fig5_radar.png")
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path

# =========================================================
# Generate all figures
# =========================================================
print("Generating figures...")
fig1_path = fig1_overview()
print(f"  Fig 1: {fig1_path}")
fig2_path = fig2_four_combinations()
print(f"  Fig 2: {fig2_path}")
fig3_path = fig3_synergy_matrix()
print(f"  Fig 3: {fig3_path}")
fig4_path = fig4_pipeline()
print(f"  Fig 4: {fig4_path}")
fig5_path = fig5_radar()
print(f"  Fig 5: {fig5_path}")

# =========================================================
# Generate DOCX
# =========================================================
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("LINKO × IONE × KOTHA\n手法組み合わせ報告書")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("3つの分析フレームワークの4通りの組み合わせとその可能性")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("2026年3月24日")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================
# Table of Contents
# ============================
doc.add_heading("目次", level=1)
toc_items = [
    "1. 背景：3つのフレームワーク",
    "  1.1 LINKO（Latent Information Normalization for Key Outcomes）",
    "  1.2 IONE（Incoherence-Oriented Neutralization and Extraction）",
    "  1.3 KOTHA（Knowledge-driven Observational-Trial Harmonization Approach）",
    "2. 4通りの組み合わせ",
    "  2.1 LINKO + IONE：ICR誘導型部分集団解析",
    "  2.2 LINKO + KOTHA：情報量を考慮したエビデンス統合",
    "  2.3 IONE + KOTHA：集団構造を反映した統合評価",
    "  2.4 LINKO + IONE + KOTHA：包括的エビデンス評価パイプライン",
    "3. シナジー比較",
    "4. 統合パイプラインの全体像",
    "5. まとめと今後の展望",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("  "):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ============================
# Section 1: Background
# ============================
doc.add_heading("1. 背景：3つのフレームワーク", level=1)

p = doc.add_paragraph()
p.add_run("本報告書では、最近開発された3つの方法論的フレームワーク——").font.size = Pt(11)
run = p.add_run("LINKO")
run.bold = True
run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
p.add_run("、")
run = p.add_run("IONE")
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
p.add_run("、")
run = p.add_run("KOTHA")
run.bold = True
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
p.add_run("——を組み合わせることで可能になる新たな分析アプローチについて報告する。これら3つのフレームワークはそれぞれ異なる分析レベル（メタ解析・個別研究・エビデンス統合）を対象としており、組み合わせは4通り存在する。")

# Figure 1
doc.add_picture(fig1_path, width=Inches(6.2))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("図1. 3つの分析フレームワークの概要と対象レベル")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 1.1 LINKO
doc.add_heading("1.1 LINKO（Latent Information Normalization for Key Outcomes）", level=2)
p = doc.add_paragraph()
p.add_run("LINKO").bold = True
p.add_run("は、メタ解析の妥当性を評価するための新しい診断的枠組みである。中核概念は")
run = p.add_run("ICR（Information Contribution Ratio：情報寄与率）")
run.bold = True
p.add_run("であり、各RCTにおいてエンドポイント変数がデータ全体の情報量のうちどれだけを占めるかを定量化する。")

doc.add_paragraph(
    "ICRには2つの算出方法がある：", style='List Bullet'
)
doc.add_paragraph(
    "ICR_v（分散ベース）：論文のTable 1から算出可能。各変数の分散比として計算", style='List Bullet'
)
doc.add_paragraph(
    "ICR_pca（主成分ベース）：個票データ（IPD）が利用可能な場合に、PCAに基づき算出", style='List Bullet'
)

p = doc.add_paragraph()
p.add_run("LINKOの独自の可視化ツールとして")
run = p.add_run("Prism Forest Plot")
run.bold = True
run.font.italic = True
p.add_run("がある。これはプリズムが光をスペクトルに分解するように、標準的なフォレストプロットにICR次元を追加し、各研究の信頼区間バーの色でICR値を、マーカーサイズで副次的ICR指標をエンコードする。また、ICR誘導型の研究選択による早期収束解析も提案されている。")

# 1.2 IONE
doc.add_heading("1.2 IONE（Incoherence-Oriented Neutralization and Extraction）", level=2)
p = doc.add_paragraph()
p.add_run("IONE").bold = True
p.add_run("は、観察研究データにおける隠れた集団構造（インコヒーレントな集団）を検出し、コヒーレントな部分集団に分解するための枠組みである。")

p = doc.add_paragraph()
p.add_run("シンプソンのパラドックスの本質は、「複数の独立してコヒーレントな集団を内包する、インコヒーレントな集団全体から得られた知見を、個々のコヒーレント集団に適用する際の危険」である。IONEはこの問題を以下のアプローチで解決する：")

doc.add_paragraph(
    "一般変数（血液検査値等）のみからPCA・クラスタリングによる特徴量空間を構成", style='List Bullet'
)
doc.add_paragraph(
    "決定力ベースおよび特徴量得点ベースの手法により、未測定の重大変数（年齢・性別等）の代理指標を構築", style='List Bullet'
)
doc.add_paragraph(
    "C1コヒーレンス指標により集団のインコヒーレンス度を定量評価", style='List Bullet'
)
doc.add_paragraph(
    "コヒーレントな部分集団の抽出により、各部分集団内で正しい知見を得る", style='List Bullet'
)

p = doc.add_paragraph()
p.add_run("IONEは交絡バイアスだけでなく、効果修飾の見落とし、生態学的誤謬、非崩壊性の問題にも対処しうる。5つの既報シンプソンのパラドックス事例（腎結石、バークレー入学、COVID-19致死率、イスラエルワクチン、喫煙・死亡率）への適用で、全事例でC1指標によるインコヒーレンス検出に成功している。")

# 1.3 KOTHA
doc.add_heading("1.3 KOTHA（Knowledge-driven Observational-Trial Harmonization Approach）", level=2)
p = doc.add_paragraph()
p.add_run("KOTHA").bold = True
p.add_run("は、観察研究とRCTの間のエビデンス乖離を診断・解決するための3モジュール構成のフレームワークである。「RCTで有意差が出ない」場合に、それが「効果がない」のか「情報量が不足している」のかを区別することを目的とする。")

# KOTHA modules table
table = doc.add_table(rows=4, cols=3)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["モジュール", "名称", "機能"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

modules = [
    ("Module K", "Counterfactual Power Simulation\n（反実仮想パワーシミュレーション）",
     "後ろ向きデータを用いて「もしRCTが実臨床のリスク分布だったら」のパワーを推定"),
    ("Module T", "Trial-Observational Bayesian Integration\n（階層ベイズ統合）",
     "RCTと観察研究を階層ベイズモデルで統合。観察研究にバイアス割引を適用"),
    ("Module H", "Hermeneutic Guideline Interpreter\n（解釈的ガイドライン指針）",
     "OIS・TSA・GRADEに基づく構造化チェックリストで低情報メタ解析を適切に解釈"),
]
for i, (mod, name, func) in enumerate(modules):
    table.rows[i+1].cells[0].text = mod
    table.rows[i+1].cells[1].text = name
    table.rows[i+1].cells[2].text = func

p = doc.add_paragraph()
p.add_run("\n急性心筋梗塞におけるマグネシウム静注療法（12試験）および心不全におけるスタチン（5観察研究+2 RCT）の実データ検証により、KOTHAが構造的情報喪失を同定し、標準的なGRADE評価よりもニュアンスのあるエビデンス評価を提供できることが示されている。")

doc.add_page_break()

# ============================
# Section 2: Four Combinations
# ============================
doc.add_heading("2. 4通りの組み合わせ", level=1)

p = doc.add_paragraph()
p.add_run("3つのフレームワークから生じる組み合わせは以下の4通りである（2つずつの組み合わせ3通り＋3つ全ての組み合わせ1通り）。各組み合わせは、単独使用では実現できない新たな分析能力を生み出す。")

# Figure 2
doc.add_picture(fig2_path, width=Inches(6.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("図2. 4通りの組み合わせとそれぞれのステップ")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 2.1 LINKO + IONE
doc.add_heading("2.1 LINKO + IONE：ICR誘導型部分集団解析", level=2)

p = doc.add_paragraph()
run = p.add_run("概要：")
run.bold = True
p.add_run("LINKOがメタ解析レベルでICRの異常を検出し、IONEが各研究内の集団構造を分解することで、メタ解析の異質性の根本原因を特定する。")

p = doc.add_paragraph()
run = p.add_run("実現される新能力：")
run.bold = True

doc.add_paragraph(
    "部分集団レベルのICR分析：IONEで抽出された各コヒーレント集団について個別にICRを算出することで、集団構造の違いがICR差異を生んでいるかを検証できる", style='List Bullet'
)
doc.add_paragraph(
    "異質性の根因特定：メタ解析でICRDが大きい場合、それが「変数次元数の違い」によるものか「集団のインコヒーレンス」によるものかを区別できる", style='List Bullet'
)
doc.add_paragraph(
    "部分集団特異的ICR加重メタ解析：コヒーレントな部分集団ごとにICR加重メタ解析を実施し、均質な集団内での統合を実現", style='List Bullet'
)
doc.add_paragraph(
    "C1-ICR複合診断：C1指標（集団コヒーレンス）とICRD（情報寄与率の不均一性）を2軸としたプロットで、各研究の問題を一目で診断", style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run("具体的シナリオ：")
run.bold = True
p.add_run("糖尿病厳格血糖コントロールのメタ解析でLINKOがUKPDS（ICR=0.125）とACCORD（ICR=0.077）の大きなICR差を検出。IONEを各試験に適用すると、UKPDSでは比較的均質な集団（C1低値）だがACCORDでは高リスク群と低リスク群の混在（C1高値）が判明。ACCORDの高リスク部分集団のみで再算出したICRはUKPDSに近づき、ICRDの主因が集団構造にあることが示される。")

# 2.2 LINKO + KOTHA
doc.add_heading("2.2 LINKO + KOTHA：情報量を考慮したエビデンス統合", level=2)

p = doc.add_paragraph()
run = p.add_run("概要：")
run.bold = True
p.add_run("LINKOが各研究の情報構造を定量化し、KOTHAのエビデンス統合にその情報を活用することで、より精密なエビデンス評価を実現する。")

p = doc.add_paragraph()
run = p.add_run("実現される新能力：")
run.bold = True

doc.add_paragraph(
    "ICR情報付きパワーシミュレーション（Module K強化）：従来のModule Kはイベント率とサンプルサイズのみに基づいていたが、ICRを組み込むことで「エンドポイントがデータ全体の何%を代表しているか」も考慮した、より現実的なパワー推定が可能になる", style='List Bullet'
)
doc.add_paragraph(
    "ICR加重ベイズ統合（Module T強化）：Module Tの階層ベイズモデルにおいて、ICRを追加の重み（またはバイアス調整項）として組み込む。ICRが高い研究ほど「エンドポイントが情報を十分に代表している」として重みを増やし、ICRが低い研究は割引する", style='List Bullet'
)
doc.add_paragraph(
    "Prism Forest Plot + TSA統合可視化：LINKOのPrism Forest PlotにKOTHAのTSA境界線を重ねることで、「情報の質（ICR）」と「情報の量（累積情報分画）」を同時に可視化", style='List Bullet'
)
doc.add_paragraph(
    "GRADE拡張（Module H強化）：Module HのチェックリストにICRDを追加のドメインとして組み込み、「情報構造の均一性」を不精確性・非直接性と並ぶ評価軸とする", style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run("具体的シナリオ：")
run.bold = True
p.add_run("スタチン療法のメタ解析において、LINKOが全研究で均一なICR（ICRD=0.009）を確認。KOTHAのModule Kでのパワーシミュレーションにこの均一性情報を加えると、「ICRが均一であるにも関わらずパワー不足」という、より強い根拠のある情報量不足の診断が可能になる。逆にICRDが大きい場合は、パワー不足とICR不均一性の両面からの診断となり、Module Hでの推奨表現もより精密になる。")

# 2.3 IONE + KOTHA
doc.add_heading("2.3 IONE + KOTHA：集団構造を反映した統合評価", level=2)

p = doc.add_paragraph()
run = p.add_run("概要：")
run.bold = True
p.add_run("IONEが観察研究の集団構造を明らかにし、KOTHAがその情報を反実仮想シミュレーションとベイズ統合に活用することで、外的妥当性の高いエビデンス評価を実現する。")

p = doc.add_paragraph()
run = p.add_run("実現される新能力：")
run.bold = True

doc.add_paragraph(
    "部分集団別反実仮想シミュレーション（Module K強化）：IONEで同定された各コヒーレント集団のリスクプロファイルを用いて、RCTが各部分集団を代表していたらどうなったかをシミュレーション。「RCTは低リスク集団のみを代表しており、高リスク集団についてはパワーゼロ」のような精密な診断が可能", style='List Bullet'
)
doc.add_paragraph(
    "部分集団特異的ベイズ統合（Module T強化）：IONEで分離した部分集団ごとに、観察研究の効果推定をRCTと統合。高リスク群での観察研究効果と低リスク群でのRCT効果を別々にモデル化し、「この治療は高リスク群でのみ有効（P(HR<1)=0.92）だが、既存RCTには高リスク群がほとんど含まれない」のような結論が可能", style='List Bullet'
)
doc.add_paragraph(
    "C1指標によるGRADE非直接性の定量化（Module H強化）：IONEのC1指標をGRADEの非直接性ドメインに組み込む。RCT対象集団のC1が低い（均質）が観察研究コホートのC1が高い（不均質）場合、「RCTは対象集団の一部のみを代表」と定量的に判定", style='List Bullet'
)
doc.add_paragraph(
    "未測定交絡の間接的評価：IONEが一般変数から重大変数の代理指標を構築できることを利用し、RCTで未測定の変数がバイアスを生んでいる可能性をKOTHAのModule Tのバイアス項に定量的に反映", style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run("具体的シナリオ：")
run.bold = True
p.add_run("急性心筋梗塞に対するマグネシウム療法で、観察研究はHR=0.54（有効）、RCTメタ解析はHR=0.56（有意ではあるがISIS-4含むとCI拡大）。IONEを観察研究コホートに適用すると、高リスク群（高齢、併存症多数）と低リスク群に分離。KOTHAのModule Kが「ISIS-4の対象は低リスク群に偏っている」と定量化し、Module Tが高リスク群のみの統合効果P(HR<1)=0.91を算出。Module Hは「高リスク群では条件付き推奨が妥当」と判定する。")

# 2.4 LINKO + IONE + KOTHA
doc.add_heading("2.4 LINKO + IONE + KOTHA：包括的エビデンス評価パイプライン", level=2)

p = doc.add_paragraph()
run = p.add_run("概要：")
run.bold = True
p.add_run("3つ全てのフレームワークを統合することで、集団構造の異質性検出→情報構造の定量化→エビデンスの統合調和という、エビデンス評価の全段階をカバーする包括的パイプラインが実現する。")

p = doc.add_paragraph()
run = p.add_run("実現される新能力：")
run.bold = True

doc.add_paragraph(
    "三次元エビデンスプロファイル：各研究を「C1（集団コヒーレンス）× ICR（情報寄与率）× 効果量」の三次元空間にプロットし、メタ解析に含まれる全研究の構造を一覧で把握", style='List Bullet'
)
doc.add_paragraph(
    "多層的品質評価：(1) IONE：研究内集団の均質性評価、(2) LINKO：エンドポイントの情報代表性評価、(3) KOTHA：エビデンス全体の統合妥当性評価——の三層で研究品質を総合評価", style='List Bullet'
)
doc.add_paragraph(
    "適応的エビデンスアップデート：新しいRCTが追加される際、IONEで集団構造を確認→LINKOでICRを算出→KOTHAで統合。「この新しいRCTはICRが低く（0.03）、かつ集団がインコヒーレント（C1=0.8）なので、重みを大幅に割引すべき」のような意思決定が可能", style='List Bullet'
)
doc.add_paragraph(
    "因果推論の強化：IONEが集団内交絡を制御し、LINKOが情報量の偏りを補正し、KOTHAがRCTと観察研究の差異を調和——三段階で因果推定の精度を段階的に向上させる", style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run("フルパイプラインの流れ：")
run.bold = True

steps_full = [
    "Phase 1（IONE）：対象となる観察研究・RCTの各データセットにIONEを適用し、コヒーレントな部分集団を同定。C1指標で各研究の集団の均質性を評価。",
    "Phase 2（LINKO）：Phase 1で同定された部分集団ごとにICRを算出。部分集団レベルでのICRDを計算し、情報構造の不均一性を評価。Prism Forest Plotで可視化。",
    "Phase 3（KOTHA）：Module Kが部分集団別のリスクプロファイル（Phase 1）とICR情報（Phase 2）を組み込んだ反実仮想パワーシミュレーションを実施。Module TがICR加重ベイズ統合を実行。Module Hが三層評価に基づく推奨を策定。",
    "Output：部分集団特異的な治療推奨、情報構造を反映したエビデンスプロファイル、追加研究のデザイン提言（どの部分集団の、どのICR水準の研究が最も必要か）。",
]
for step in steps_full:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# ============================
# Section 3: Synergy Comparison
# ============================
doc.add_heading("3. シナジー比較", level=1)

p = doc.add_paragraph()
p.add_run("以下に各組み合わせの能力プロファイルを示す。6つの評価軸（異質性診断・集団分解・情報定量化・エビデンス統合・ガイドライン有用性・外的妥当性）について、5段階で評価した。")

# Figure 3
doc.add_picture(fig3_path, width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("図3. シナジーマトリックス：各組み合わせが実現する新能力と利得")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Figure 5 (radar)
doc.add_picture(fig5_path, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("図4. レーダーチャート：各組み合わせの能力プロファイル比較")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.add_run("レーダーチャートから明らかなように、2つの組み合わせはそれぞれ特定の軸に強みを持つ一方で弱点も残る。3つ全ての統合のみが全軸で最大評価を達成する。")

# Comparison table
doc.add_heading("組み合わせ別の特徴比較", level=3)
comp_table = doc.add_table(rows=5, cols=4)
comp_table.style = 'Medium Shading 1 Accent 1'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

comp_headers = ["組み合わせ", "主な強み", "主な弱点", "最適な適用場面"]
for i, h in enumerate(comp_headers):
    cell = comp_table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

comp_data = [
    ("LINKO + IONE", "異質性の根因特定", "エビデンス統合力が弱い", "メタ解析の異質性が高く原因不明の場合"),
    ("LINKO + KOTHA", "情報構造を反映した精密統合", "集団構造を無視", "ICRが研究間で異なるメタ解析の再評価"),
    ("IONE + KOTHA", "外的妥当性の高い統合評価", "情報構造の定量化が不十分", "RCTと観察研究の乖離が集団差に起因する場合"),
    ("全3つ統合", "全評価軸で最大能力", "実装の複雑さとデータ要件", "包括的なエビデンス再評価が必要な場合"),
]
for i, (combo, strength, weakness, use_case) in enumerate(comp_data):
    comp_table.rows[i+1].cells[0].text = combo
    comp_table.rows[i+1].cells[1].text = strength
    comp_table.rows[i+1].cells[2].text = weakness
    comp_table.rows[i+1].cells[3].text = use_case

doc.add_page_break()

# ============================
# Section 4: Pipeline
# ============================
doc.add_heading("4. 統合パイプラインの全体像", level=1)

p = doc.add_paragraph()
p.add_run("3つのフレームワークを統合したフルパイプラインの全体像を図5に示す。データの流れ（実線矢印）に加えて、モジュール間のフィードバック（破線矢印）も示している。")

# Figure 4
doc.add_picture(fig4_path, width=Inches(6.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("図5. 統合パイプラインの全体像：データフロー（実線）とフィードバック（破線）")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.add_run("このパイプラインでは、各フレームワークの出力が他のフレームワークの入力として利用される。特に重要なフィードバックパスは以下の2つである：")

doc.add_paragraph(
    "IONE → KOTHA Module K：IONEで同定された部分集団のリスクプロファイルが、Module Kの反実仮想パワーシミュレーションの入力となる。これにより「RCTがどの部分集団を見落としているか」を定量的に評価できる。", style='List Bullet'
)
doc.add_paragraph(
    "LINKO → KOTHA Module T：LINKOで算出されたICRが、Module Tの階層ベイズモデルにおける追加の重み情報となる。ICRが高い研究はエンドポイントの情報代表性が高いため、ベイズ統合でより大きな影響力を持つ。", style='List Bullet'
)

doc.add_page_break()

# ============================
# Section 5: Summary
# ============================
doc.add_heading("5. まとめと今後の展望", level=1)

p = doc.add_paragraph()
p.add_run("まとめ").bold = True
doc.add_paragraph(
    "LINKO、IONE、KOTHAはそれぞれ異なる分析レベルを対象とするが、組み合わせることで単独では実現できないシナジーが生まれる。", style='List Bullet'
)
doc.add_paragraph(
    "4通りの組み合わせはそれぞれ固有の強みを持ち、分析目的に応じた選択が可能。", style='List Bullet'
)
doc.add_paragraph(
    "3つ全ての統合パイプラインは、エビデンス評価の全段階（集団構造→情報構造→統合）をカバーする唯一のアプローチである。", style='List Bullet'
)
doc.add_paragraph(
    "いずれの組み合わせも、既存のメタ解析やGRADEフレームワークと補完的に使用でき、エビデンスの透明性と解釈可能性を向上させる。", style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run("\n今後の展望")
run.bold = True

doc.add_paragraph(
    "統合パイプラインのソフトウェア実装（Python/Rパッケージとして公開）", style='List Bullet'
)
doc.add_paragraph(
    "大規模検証：Cochrane Database of Systematic Reviewsの複数レビューへの適用", style='List Bullet'
)
doc.add_paragraph(
    "ICR-C1複合指標の形式的定義と閾値の確立", style='List Bullet'
)
doc.add_paragraph(
    "ガイドライン作成委員会での実用性評価（Module Hへの組み込み）", style='List Bullet'
)
doc.add_paragraph(
    "個票データメタ解析（IPD-MA）における統合パイプラインの完全実装と検証", style='List Bullet'
)

# Save
output_path = "/home/ubuntu/LINKO_IONE_KOTHA_combinations_report.docx"
doc.save(output_path)
print(f"\nReport saved to: {output_path}")
